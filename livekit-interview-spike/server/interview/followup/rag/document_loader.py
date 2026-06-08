from __future__ import annotations

import hashlib
import re
from pathlib import Path

from pypdf import PdfReader

from interview.followup.rag.models import MetadataValue, MethodologySource


SUPPORTED_DOCUMENT_SUFFIXES = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def load_methodology_documents(
    *,
    path: str | Path,
    source_type: str = "internal_note",
    license_note: str = "",
    metadata: dict[str, MetadataValue] | None = None,
    include_index_documents: bool = False,
) -> list[MethodologySource]:
    root = Path(path)
    files = _document_files(root)
    sources: list[MethodologySource] = []
    for file in files:
        if not include_index_documents and _is_index_document(file):
            continue
        document = _extract_document(file)
        if not document.content.strip():
            continue
        merged_metadata = {**document.metadata, **dict(metadata or {})}
        sources.append(
            MethodologySource(
                source_id=_source_id(file),
                title=document.title or file.stem,
                content=document.content,
                source_type=str(document.metadata.get("source_type") or source_type),
                source_url=str(document.metadata.get("source_file") or file),
                license_note=str(document.metadata.get("license_note") or license_note),
                metadata=merged_metadata,
            )
        )
    return sources


def _document_files(path: Path) -> list[Path]:
    if path.is_file():
        if path.suffix.lower() not in SUPPORTED_DOCUMENT_SUFFIXES:
            raise ValueError(f"Unsupported document suffix: {path.suffix}")
        return [path]
    if not path.exists():
        raise FileNotFoundError(path)
    return sorted(
        file
        for file in path.rglob("*")
        if file.is_file() and file.suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES
    )


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    raise ValueError(f"Unsupported document suffix: {suffix}")


class _ExtractedDocument:
    def __init__(self, *, title: str, content: str, metadata: dict[str, MetadataValue]) -> None:
        self.title = title
        self.content = content
        self.metadata = metadata


def _extract_document(path: Path) -> _ExtractedDocument:
    text = _extract_text(path)
    metadata: dict[str, MetadataValue] = {}
    if path.suffix.lower() in {".md", ".markdown"}:
        text, metadata = _strip_frontmatter(text)
        text = _strip_markdown_noise(text)
    return _ExtractedDocument(
        title=str(metadata.get("title") or _first_markdown_heading(text) or path.stem),
        content=text,
        metadata=metadata,
    )


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = [page.extract_text() or "" for page in reader.pages]
    return _normalize_text("\n".join(parts))


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)
    return _normalize_text("\n".join([*paragraphs, *table_cells]))


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_frontmatter(text: str) -> tuple[str, dict[str, MetadataValue]]:
    if not text.startswith("---\n"):
        return text, {}
    end = text.find("\n---", 4)
    if end < 0:
        return text, {}
    raw_metadata = text[4:end].strip()
    content = text[end + len("\n---") :].lstrip()
    return content, _parse_simple_frontmatter(raw_metadata)


def _parse_simple_frontmatter(raw_metadata: str) -> dict[str, MetadataValue]:
    metadata: dict[str, MetadataValue] = {}
    current_list_key = ""
    for line in raw_metadata.splitlines():
        stripped = line.strip()
        if current_list_key and stripped.startswith("- "):
            value = stripped[2:].strip().strip('"').strip("'")
            existing = metadata.setdefault(current_list_key, [])
            if isinstance(existing, list) and value:
                existing.append(value)
            continue
        current_list_key = ""
        if ":" not in line:
            continue
        key, raw_value = line.split(":", 1)
        key = key.strip()
        value = raw_value.strip().strip('"').strip("'")
        if value.startswith("[") and value.endswith("]"):
            metadata[key] = [item.strip().strip('"').strip("'") for item in value[1:-1].split(",") if item.strip()]
        elif not value:
            metadata[key] = []
            current_list_key = key
        elif value.lower() in {"true", "false"}:
            metadata[key] = value.lower() == "true"
        else:
            metadata[key] = _coerce_metadata_value(key, value)
    return metadata


def _strip_markdown_noise(text: str) -> str:
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"```(?:yaml|yml)\s+.*?chunk_id:.*?```\s*", "", text, flags=re.DOTALL | re.IGNORECASE)
    cleaned_lines: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if "scripts\\ingest_documents.py" in stripped or "scripts/ingest_documents.py" in stripped:
            continue
        if stripped.startswith("> 用途：") or stripped.startswith("> 用途:"):
            continue
        cleaned_lines.append(line)
    return _normalize_text("\n".join(cleaned_lines))


def _is_index_document(path: Path) -> bool:
    stem = path.stem.lower()
    return stem in {"index", "readme"} or "index" in stem


def _first_markdown_heading(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    return ""


def _coerce_metadata_value(key: str, value: str) -> MetadataValue:
    list_keys = {
        "strategy_categories",
        "role_families",
        "competency_archetypes",
        "evidence_categories",
        "tags",
    }
    if key in list_keys and " " in value and "," not in value:
        return [item.strip() for item in value.split() if item.strip()]
    return value


def _source_id(path: Path) -> str:
    resolved = str(path.resolve()).encode("utf-8", errors="ignore")
    digest = hashlib.sha1(resolved).hexdigest()[:12]
    safe_stem = re.sub(r"[^a-zA-Z0-9_-]+", "-", path.stem).strip("-").lower() or "document"
    return f"doc-{safe_stem}-{digest}"
